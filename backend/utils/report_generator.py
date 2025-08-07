"""报告生成工具
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os
from config.config import FILE_PATHS


class ReportGenerator:
    """报告生成器
    
    用于生成标准化格式的Word版本报告
    """
    
    def __init__(self):
        """初始化报告生成器
        """
        pass
    
    def generate_hci_nbti_report(self, data, start_date=None, end_date=None):
        """生成HCI&NBTI分析报告
        
        Args:
            data (list): HCI&NBTI数据
            start_date (str): 开始日期
            end_date (str): 结束日期
            
        Returns:
            str: 生成的报告文件路径
        """
        # 创建文档
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加标题
        title = doc.add_heading('HCI&NBTI数据分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加报告信息
        doc.add_paragraph(f'报告生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        if start_date and end_date:
            doc.add_paragraph(f'数据范围: {start_date} 至 {end_date}')
        doc.add_paragraph(f'数据记录数: {len(data)}')
        
        # 添加数据表格
        if data:
            doc.add_heading('数据详情', level=1)
            
            # 创建表格
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # 设置表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Lot ID'
            hdr_cells[2].text = '设备ID'
            hdr_cells[3].text = '设备名称'
            hdr_cells[4].text = 'STD'
            hdr_cells[5].text = 'VTD'
            
            # 填充数据
            for item in data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('id', ''))
                row_cells[1].text = str(item.get('lot_id', ''))
                row_cells[2].text = str(item.get('device_id', ''))
                row_cells[3].text = str(item.get('device_name', ''))
                row_cells[4].text = str(item.get('std', ''))
                row_cells[5].text = str(item.get('vtd', ''))
        
        # 生成文件名
        filename = f"hci_nbti_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(FILE_PATHS['reports'], filename)
        
        # 确保报告目录存在
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # 保存文档
        doc.save(file_path)
        
        return file_path
    
    def generate_vrdb_report(self, data, start_date=None, end_date=None):
        """生成VRDB分析报告
        
        Args:
            data (list): VRDB数据
            start_date (str): 开始日期
            end_date (str): 结束日期
            
        Returns:
            str: 生成的报告文件路径
        """
        # 创建文档
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加标题
        title = doc.add_heading('VRDB数据分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加报告信息
        doc.add_paragraph(f'报告生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        if start_date and end_date:
            doc.add_paragraph(f'数据范围: {start_date} 至 {end_date}')
        doc.add_paragraph(f'数据记录数: {len(data)}')
        
        # 添加数据表格
        if data:
            doc.add_heading('数据详情', level=1)
            
            # 创建表格
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # 设置表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = '设备ID'
            hdr_cells[2].text = '测试日期'
            hdr_cells[3].text = '电压(V)'
            hdr_cells[4].text = '电流(A)'
            hdr_cells[5].text = '电阻(Ω)'
            
            # 填充数据
            for item in data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('id', ''))
                row_cells[1].text = str(item.get('device_id', ''))
                row_cells[2].text = str(item.get('test_date', ''))
                row_cells[3].text = str(item.get('voltage', ''))
                row_cells[4].text = str(item.get('current', ''))
                row_cells[5].text = str(item.get('resistance', ''))
        
        # 生成文件名
        filename = f"vrdb_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(FILE_PATHS['reports'], filename)
        
        # 确保报告目录存在
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # 保存文档
        doc.save(file_path)
        
        return file_path
    
    def generate_tddb_report(self, data, start_date=None, end_date=None):
        """生成TDDB分析报告
        
        Args:
            data (list): TDDB数据
            start_date (str): 开始日期
            end_date (str): 结束日期
            
        Returns:
            str: 生成的报告文件路径
        """
        # 创建文档
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加标题
        title = doc.add_heading('TDDB数据分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加报告信息
        doc.add_paragraph(f'报告生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        if start_date and end_date:
            doc.add_paragraph(f'数据范围: {start_date} 至 {end_date}')
        doc.add_paragraph(f'数据记录数: {len(data)}')
        
        # 添加数据表格
        if data:
            doc.add_heading('数据详情', level=1)
            
            # 创建表格
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            
            # 设置表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = '设备ID'
            hdr_cells[2].text = 'Wafer ID'
            hdr_cells[3].text = 'X坐标'
            hdr_cells[4].text = 'Y坐标'
            hdr_cells[5].text = '数值'
            hdr_cells[6].text = '测试日期'
            
            # 填充数据
            for item in data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('id', ''))
                row_cells[1].text = str(item.get('device_id', ''))
                row_cells[2].text = str(item.get('wafer_id', ''))
                row_cells[3].text = str(item.get('x_coordinate', ''))
                row_cells[4].text = str(item.get('y_coordinate', ''))
                row_cells[5].text = str(item.get('value', ''))
                row_cells[6].text = str(item.get('test_date', ''))
        
        # 生成文件名
        filename = f"tddb_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(FILE_PATHS['reports'], filename)
        
        # 确保报告目录存在
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # 保存文档
        doc.save(file_path)
        
        return file_path
    
    def generate_comprehensive_report(self, hci_nbti_data, vrdb_data, tddb_data, start_date=None, end_date=None):
        """生成综合分析报告
        
        Args:
            hci_nbti_data (list): HCI&NBTI数据
            vrdb_data (list): VRDB数据
            tddb_data (list): TDDB数据
            start_date (str): 开始日期
            end_date (str): 结束日期
            
        Returns:
            str: 生成的报告文件路径
        """
        # 创建文档
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加标题
        title = doc.add_heading('制程可靠性综合分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加报告信息
        doc.add_paragraph(f'报告生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        if start_date and end_date:
            doc.add_paragraph(f'数据范围: {start_date} 至 {end_date}')
        
        # 添加HCI&NBTI数据分析部分
        doc.add_heading('1. HCI&NBTI数据分析', level=1)
        doc.add_paragraph(f'数据记录数: {len(hci_nbti_data)}')
        
        # 添加VRDB数据分析部分
        doc.add_heading('2. VRDB数据分析', level=1)
        doc.add_paragraph(f'数据记录数: {len(vrdb_data)}')
        
        # 添加TDDB数据分析部分
        doc.add_heading('3. TDDB数据分析', level=1)
        doc.add_paragraph(f'数据记录数: {len(tddb_data)}')
        
        # 生成文件名
        filename = f"comprehensive_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(FILE_PATHS['reports'], filename)
        
        # 确保报告目录存在
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # 保存文档
        doc.save(file_path)
        
        return file_path